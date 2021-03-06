#get_file_word.py
#!/usr/bin/env python
# coding: utf-8

# In[ ]:


# Importar bibliotecas
from os import listdir
from os.path import isfile, join
import docx2txt
import docx
import os
import glob

from os import listdir
from os.path import isfile, join

#argumentos 
#fileDir = r"/Users/wpessoa/repositorios/Fluxo_Editorial/words/"
#fileExt = r".docx"

def get_info(fileDir,files_array):
    while True:
        try:
            #print(os.path.exists(fileDir, fileDir)) 
            #files_array = [_ for _ in os.listdir(fileDir) if _.endswith(fileExt)]
            indice = 0
            image_array=[]
            table_chunks = []
            table_return = []
            
            for file in files_array:
              
                # abrir conecção com o arquivo Word 
                #doc = docx.Document(fileDir+'/'+file)
                doc = docx.Document(fileDir+file)


                # data criação
                dt_doc = doc.core_properties.created
                # ler em cada parágrafo dentro do arquivo Word
                paragra= [p.text for p in doc.paragraphs]

                # extai a quantidade de caracteres
                #caracteres = docx2txt.process(fileDir+'/'+file)
                caracteres = docx2txt.process(fileDir+file)


                # gerar um array com tamanho de todas as imagens
                for image in doc.inline_shapes:
                    image_array.append([image.width, image.height])
                    # print (image.width, image.height)

                # gerar um array com todas as tabelas
                for table in doc.tables:
                    table_chunks.append(table)
                    
                # gerar lista resposta
                list_retur=file[2:8],len(caracteres), len(table_chunks), len(image_array),dt_doc.date()
                
                # gerar tabela resposta
                table_return.append(list_retur)
            #return(files_array[indice][2:8],len(caracteres), len(table_chunks), len(image_array),dt_doc.date())
            return table_return
            #break
        except ValueError:
            print("Oops!  erro no método verificar diret'ório e extensão do arquivos no argumento desse método")
           

    

